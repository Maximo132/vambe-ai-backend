import logging
import os
import io
import magic
import PyPDF2
import docx
import uuid
from datetime import datetime
from typing import List, Optional, Dict, Any, BinaryIO, Union, AsyncGenerator
from pathlib import Path
from fastapi import UploadFile, HTTPException, status
from sqlalchemy.orm import Session, selectinload
from sqlalchemy import select, and_, or_

from app.core.config import settings
from app.models.document import Document as DocumentModel, DocumentType, DocumentStatus
from app.models.knowledge_base import KnowledgeBase, KnowledgeDocument
from app.schemas.document import (
    DocumentBase, DocumentCreate, DocumentUpdate, DocumentInDB, DocumentProcessRequest,
    DocumentProcessResponse, DocumentSearchQuery, DocumentSearchResults, DocumentUploadResponse
)
from app.services.base import CRUDBase
from app.services.openai_service import openai_service
from app.services.weaviate_service import weaviate_service
from app.utils.file_utils import save_upload_file, get_file_extension, generate_unique_filename

logger = logging.getLogger(__name__)

class DocumentService(CRUDBase[DocumentModel, DocumentCreate, DocumentUpdate]):
    """
    Servicio para la gestión de documentos, incluyendo carga, procesamiento y búsqueda.
    """
    
    def __init__(self, db: Session):
        super().__init__(DocumentModel, db)
        self.upload_dir = Path(settings.UPLOAD_DIR) / "documents"
        self.upload_dir.mkdir(parents=True, exist_ok=True)
    
    async def upload_document(
        self,
        file: UploadFile,
        user_id: int,
        metadata: Optional[Dict[str, Any]] = None
    ) -> DocumentUploadResponse:
        """
        Sube un documento al servidor y crea un registro en la base de datos.
        
        Args:
            file: Archivo a subir
            user_id: ID del usuario que sube el documento
            metadata: Metadatos adicionales
            
        Returns:
            DocumentUploadResponse con los detalles del documento subido
        """
        try:
            # Validar el tipo de archivo
            file_type = self._detect_file_type(file.filename, await file.read(1024), file.content_type)
            await file.seek(0)  # Volver al inicio del archivo
            
            # Generar un nombre de archivo único
            file_ext = get_file_extension(file.filename)
            filename = generate_unique_filename(file.filename, prefix=f"doc_{user_id}_")
            file_path = self.upload_dir / filename
            
            # Guardar el archivo
            await save_upload_file(file, file_path)
            
            # Crear el registro en la base de datos
            document_data = DocumentCreate(
                title=file.filename,
                description=metadata.get("description") if metadata else None,
                file_url=str(file_path),
                file_type=DocumentType(file_type),
                metadata=metadata or {}
            )
            
            db_document = await self.create(document_data, user_id=user_id)
            
            return DocumentUploadResponse(
                **db_document.dict(),
                message="Documento cargado exitosamente"
            )
            
        except Exception as e:
            logger.error(f"Error al subir documento: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Error al subir el documento: {str(e)}"
            )
    
    async def process_document(
        self,
        document_id: int,
        process_data: DocumentProcessRequest,
        user_id: int
    ) -> DocumentProcessResponse:
        """
        Procesa un documento para extraer texto, dividirlo en fragmentos y generar embeddings.
        
        Este método se encarga de:
        1. Verificar la existencia y permisos del documento
        2. Extraer el texto del documento según su tipo
        3. Dividir el texto en fragmentos manejables
        4. Generar embeddings para cada fragmento
        5. Almacenar los fragmentos en Weaviate
        6. Actualizar el estado del documento
        
        Args:
            document_id: ID del documento a procesar
            process_data: Parámetros de procesamiento (tamaño de fragmentos, etc.)
            user_id: ID del usuario que realiza la operación
            
        Returns:
            DocumentProcessResponse: Resultado del procesamiento con estadísticas
            
        Raises:
            HTTPException: Si el documento no existe, no tiene permisos o hay un error en el procesamiento
        """
        try:
            # Obtener el documento con bloqueo para evitar procesamiento concurrente
            document = await self.get(
                document_id=document_id, 
                user_id=user_id, 
                include_knowledge_bases=True,
                for_update=True
            )
            if not document:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail=f"Documento con ID {document_id} no encontrado o acceso denegado"
                )
                
            # Verificar que el documento no esté siendo procesado actualmente
            if document.status == DocumentStatus.PROCESSING:
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail="El documento ya está siendo procesado"
                )
                
            # Actualizar el estado del documento a PROCESANDO
            document.status = DocumentStatus.PROCESSING
            document.processing_started_at = datetime.utcnow()
            document.metadata = document.metadata or {}
            document.metadata.update({
                "processing_attempts": document.metadata.get("processing_attempts", 0) + 1,
                "last_processing_attempt": datetime.utcnow().isoformat(),
                "processing_params": process_data.dict()
            })
            
            await self.db.commit()
            await self.db.refresh(document)
            
            # Leer el contenido del archivo
            try:
                with open(document.file_url, "rb") as f:
                    content = f.read()
            except Exception as e:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"No se pudo leer el archivo del documento: {str(e)}"
                )
            
            # Extraer texto según el tipo de archivo
            try:
                text = await self._extract_text(document.file_url, content, document.doc_type)
                if not text.strip():
                    raise ValueError("No se pudo extraer texto del documento o está vacío")
            except Exception as e:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Error al extraer texto del documento: {str(e)}"
                )
            
            # Dividir el texto en fragmentos
            try:
                chunks = self._chunk_text(
                    text,
                    chunk_size=process_data.chunk_size,
                    chunk_overlap=process_data.chunk_overlap
                )
                if not chunks:
                    raise ValueError("No se pudieron generar fragmentos del texto")
            except Exception as e:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Error al dividir el texto en fragmentos: {str(e)}"
                )
            
            # Procesar cada fragmento
            chunks_with_embeddings = []
            failed_chunks = 0
            
            for i, chunk in enumerate(chunks):
                try:
                    # Generar un ID único para el fragmento
                    chunk_id = f"{document_id}_{i}"
                    
                    # Generar embedding usando el servicio Weaviate
                    embedding = await weaviate_service.get_embedding(chunk["text"])
                    
                    # Almacenar en Weaviate
                    weaviate_id = await weaviate_service.add_document_chunk(
                        document_id=str(document_id),
                        chunk_id=chunk_id,
                        text=chunk["text"],
                        metadata={
                            "document_title": document.title,
                            "document_type": document.doc_type.value if document.doc_type else None,
                            "user_id": str(user_id),
                            "chunk_index": i,
                            "chunk_chars": len(chunk["text"]),
                            **chunk.get("metadata", {}),
                            **process_data.metadata
                        },
                        vector=embedding
                    )
                    
                    chunks_with_embeddings.append({
                        "chunk_id": chunk_id,
                        "weaviate_id": weaviate_id,
                        "text_length": len(chunk["text"]),
                        **chunk.get("metadata", {})
                    })
                    
                except Exception as e:
                    failed_chunks += 1
                    logger.error(f"Error procesando fragmento {i} del documento {document_id}: {str(e)}")
            
            # Verificar si se procesaron fragmentos correctamente
            if not chunks_with_embeddings:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="No se pudo procesar ningún fragmento del documento"
                )
            
            # Actualizar el estado del documento a COMPLETADO
            document.status = DocumentStatus.COMPLETED
            document.processed_at = datetime.utcnow()
            document.metadata.update({
                "chunks_processed": len(chunks_with_embeddings),
                "chunks_failed": failed_chunks,
                "total_chunks": len(chunks),
                "processing_time_seconds": (datetime.utcnow() - document.processing_started_at).total_seconds(),
                "processing_completed_at": datetime.utcnow().isoformat()
            })
            
            # Si el documento está asociado a alguna base de conocimiento, actualizar sus metadatos
            if hasattr(document, 'knowledge_bases') and document.knowledge_bases:
                document.metadata["knowledge_bases"] = [
                    {"id": kb.knowledge_base_id, "name": kb.knowledge_base.name}
                    for kb in document.knowledge_bases
                    if kb.knowledge_base
                ]
            
            await self.db.commit()
            
            logger.info(f"Documento {document_id} procesado exitosamente. "
                       f"{len(chunks_with_embeddings)} fragmentos procesados, {failed_chunks} fallidos")
            
            return DocumentProcessResponse(
                document_id=document_id,
                status="completed",
                message=(
                    f"Documento procesado exitosamente. "
                    f"{len(chunks_with_embeddings)} fragmentos generados, {failed_chunks} fallidos."
                ),
                chunks_processed=len(chunks_with_embeddings),
                chunks_failed=failed_chunks,
                total_chunks=len(chunks)
            )
            
        except HTTPException:
            # Re-lanzar excepciones HTTP
            raise
            
        except Exception as e:
            logger.error(f"Error inesperado al procesar documento {document_id}: {str(e)}", exc_info=True)
            
            # Actualizar el estado del documento a FALLIDO
            try:
                if 'document' in locals():
                    document.status = DocumentStatus.FAILED
                    document.metadata = document.metadata or {}
                    document.metadata.update({
                        "error": str(e),
                        "error_type": e.__class__.__name__,
                        "error_timestamp": datetime.utcnow().isoformat(),
                        "processing_time_seconds": (
                            (datetime.utcnow() - document.processing_started_at).total_seconds()
                            if hasattr(document, 'processing_started_at') and document.processing_started_at
                            else None
                        )
                    })
                    await self.db.commit()
            except Exception as update_error:
                logger.error(f"Error al actualizar el estado del documento a FALLIDO: {str(update_error)}")
            
            # Lanzar excepción con el error original
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error inesperado al procesar el documento: {str(e)}"
            )
    
    async def search_documents(
        self,
        query: str,
        user_id: int,
        search_params: DocumentSearchQuery,
        knowledge_base_id: Optional[int] = None
    ) -> DocumentSearchResults:
        """
        Busca documentos relevantes para una consulta utilizando embeddings semánticos.
        
        Args:
            query: Texto de búsqueda
            user_id: ID del usuario que realiza la búsqueda
            search_params: Parámetros de búsqueda
            knowledge_base_id: ID opcional de la base de conocimiento para filtrar
            
        Returns:
            DocumentSearchResults con los resultados de la búsqueda
        """
        try:
            # Obtener el embedding de la consulta
            query_embedding = await weaviate_service.get_embedding(query)
            
            # Construir filtros de búsqueda
            filters = {"user_id": str(user_id)}
            
            # Aplicar filtros adicionales si se proporcionan
            if search_params.filters:
                filters.update(search_params.filters)
            
            # Si se especifica un knowledge_base_id, filtrar por él
            kb_id = knowledge_base_id or search_params.knowledge_base_id
            if kb_id:
                filters["knowledge_base_id"] = str(kb_id)
            
            # Realizar la búsqueda en Weaviate
            search_results = await weaviate_service.semantic_search(
                query_embedding=query_embedding,
                limit=search_params.limit,
                min_score=search_params.min_score,
                filters=filters if filters else None,
                include_vectors=search_params.include_vectors,
                include_metadata=search_params.include_metadata
            )
            
            # Si no hay resultados, devolver una lista vacía
            if not search_results:
                return DocumentSearchResults(results=[], total=0)
            
            # Obtener los IDs de los documentos encontrados
            doc_ids = [int(result["document_id"]) for result in search_results if result.get("document_id")]
            
            if not doc_ids:
                return DocumentSearchResults(results=[], total=0)
            
            # Obtener los documentos completos de la base de datos con sus relaciones
            stmt = (
                select(DocumentModel)
                .where(DocumentModel.id.in_(doc_ids))
                .options(
                    selectinload(DocumentModel.knowledge_bases)
                    .selectinload(KnowledgeDocument.knowledge_base)
                )
            )
            
            result = await self.db.execute(stmt)
            documents = result.scalars().all()
            
            # Crear un diccionario para búsqueda rápida por ID
            docs_by_id = {doc.id: doc for doc in documents}
            
            # Construir los resultados de búsqueda
            results = []
            for item in search_results:
                doc_id = item.get("document_id")
                if not doc_id or int(doc_id) not in docs_by_id:
                    continue
                    
                doc = docs_by_id[int(doc_id)]
                
                # Crear el objeto DocumentInDB con la información del documento
                doc_data = DocumentInDB.model_validate(doc)
                
                # Si se solicitan las bases de conocimiento, incluirlas
                if hasattr(doc, 'knowledge_bases') and doc.knowledge_bases:
                    doc_data.knowledge_bases = [
                        {
                            'id': kb.knowledge_base.id,
                            'name': kb.knowledge_base.name,
                            'description': kb.knowledge_base.description,
                            'is_public': kb.knowledge_base.is_public,
                            'added_at': kb.created_at.isoformat() if kb.created_at else None,
                            'metadata': kb.metadata or {}
                        }
                        for kb in doc.knowledge_bases
                    ]
                
                # Crear el resultado de búsqueda
                search_result = DocumentSearchResult(
                    document=doc_data,
                    score=item.get("score", 0.0),
                    chunk_text=item.get("text"),
                    chunk_metadata=item.get("metadata")
                )
                
                # Si se solicitan los vectores, incluirlos
                if search_params.include_vectors and "vector" in item:
                    search_result.vector = item["vector"]
                
                results.append(search_result)
            
            return DocumentSearchResults(
                results=results,
                total=len(results)
            )
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error en la búsqueda de documentos: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error al buscar documentos: {str(e)}"
            )
    
    async def _extract_text(
        self,
        file_path: str,
        content: bytes,
        file_type: DocumentType
    ) -> str:
        """
        Extrae texto de un archivo según su tipo.
        
        Args:
            file_path: Ruta del archivo (para propósitos de registro)
            content: Contenido binario del archivo
            file_type: Tipo de documento (PDF, DOCX, TXT, MARKDOWN)
            
        Returns:
            str: Texto extraído del archivo
            
        Raises:
            ValueError: Si el tipo de archivo no es compatible o hay un error en la extracción
        """
        try:
            logger.info(f"Extrayendo texto de archivo {file_path} de tipo {file_type}")
            
            if not content or len(content) == 0:
                raise ValueError("El contenido del archivo está vacío")
                
            # Convertir a string si es necesario (para compatibilidad con versiones anteriores)
            if isinstance(file_type, str):
                file_type = DocumentType(file_type.lower())
                
            # Manejar diferentes tipos de archivo
            if file_type == DocumentType.PDF:
                try:
                    return self._extract_text_from_pdf(io.BytesIO(content))
                except Exception as pdf_error:
                    logger.error(f"Error al extraer texto de PDF: {str(pdf_error)}")
                    # Intentar con un método alternativo si el primero falla
                    return self._extract_text_with_fallback(content, 'pdf')
                    
            elif file_type == DocumentType.DOCX:
                try:
                    return self._extract_text_from_docx(io.BytesIO(content))
                except Exception as docx_error:
                    logger.error(f"Error al extraer texto de DOCX: {str(docx_error)}")
                    # Intentar con un método alternativo si el primero falla
                    return self._extract_text_with_fallback(content, 'docx')
                    
            elif file_type in [DocumentType.TXT, DocumentType.MARKDOWN]:
                try:
                    # Intentar con diferentes codificaciones comunes
                    encodings = ['utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
                    for encoding in encodings:
                        try:
                            return content.decode(encoding)
                        except UnicodeDecodeError:
                            continue
                    # Si ninguna codificación funcionó, intentar con reemplazos de errores
                    return content.decode('utf-8', errors='replace')
                except Exception as txt_error:
                    logger.error(f"Error al decodificar archivo de texto: {str(txt_error)}")
                    raise ValueError(f"No se pudo decodificar el archivo de texto: {str(txt_error)}")
            else:
                raise ValueError(f"Tipo de archivo no soportado: {file_type}")
                
        except Exception as e:
            logger.error(f"Error al extraer texto de {file_path}: {str(e)}", exc_info=True)
            raise ValueError(f"Error al extraer texto del archivo: {str(e)}")
    
    def _detect_file_type(self, filename: str, content: bytes, content_type: str) -> str:
        """Detecta el tipo de archivo usando magic y la extensión."""
        try:
            # Usar magic para detectar el tipo MIME real
            mime = magic.Magic(mime=True)
            detected_mime = mime.from_buffer(content)
            
            # Mapear tipos MIME a tipos de documento
            mime_type_map = {
                'application/pdf': 'pdf',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
                'text/plain': 'txt',
                'text/markdown': 'markdown',
                'text/x-markdown': 'markdown',
            }
            
            file_type = mime_type_map.get(detected_mime.lower())
            
            # Si no se pudo detectar por MIME, intentar por extensión
            if not file_type:
                ext = filename.split('.')[-1].lower()
                ext_map = {
                    'pdf': 'pdf',
                    'docx': 'docx',
                    'txt': 'txt',
                    'md': 'markdown',
                    'markdown': 'markdown'
                }
                file_type = ext_map.get(ext)
            
            if not file_type:
                raise ValueError("Tipo de archivo no soportado")
                
            return file_type
            
        except Exception as e:
            logger.error(f"Error al detectar tipo de archivo: {str(e)}")
            raise ValueError("No se pudo determinar el tipo de archivo")
    
    def _chunk_text(
        self,
        text: str,
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ) -> List[Dict[str, Any]]:
        """
        Divide el texto en fragmentos con superposición.
        
        Args:
            text: Texto a dividir
            chunk_size: Tamaño máximo de cada fragmento en caracteres
            chunk_overlap: Número de caracteres de superposición entre fragmentos
            
        Returns:
            Lista de fragmentos con metadatos
        """
        if not text:
            return []
            
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = min(start + chunk_size, text_length)
            
            # Asegurarse de no cortar palabras a la mitad
            if end < text_length:
                # Buscar el último espacio en blanco o salto de línea
                split_at = text.rfind(' ', start, end + 1)
                if split_at == -1 or split_at < start + chunk_size // 2:
                    split_at = text.find(' ', end, end + 100)  # Buscar el siguiente espacio
                    if split_at == -1:
                        split_at = end
                end = split_at
            
            chunk_text = text[start:end].strip()
            if chunk_text:  # Solo agregar fragmentos no vacíos
                chunks.append({
                    "text": chunk_text,
                    "metadata": {
                        "chunk_start": start,
                        "chunk_end": end,
                        "chunk_length": end - start
                    }
                })
            
            # Mover el inicio al final del fragmento actual, menos el solapamiento
            start = max(start + 1, end - chunk_overlap)
            
            # Si no hay progreso, forzar el avance para evitar bucles infinitos
            if start >= end and end < text_length:
                start = end
        
        return chunks
    
    def _extract_text_from_pdf(self, file_stream: BinaryIO) -> str:
        """Extrae texto de un archivo PDF."""
        try:
            pdf_reader = PyPDF2.PdfReader(file_stream)
            text_parts = []
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text() or ""  # Manejar páginas sin texto
                text_parts.append(f"--- Página {page_num + 1} ---\n{text}")
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error al extraer texto de PDF: {str(e)}")
            raise
    
    def _extract_text_from_docx(self, file_stream: BinaryIO) -> str:
        """Extrae texto de un archivo DOCX."""
        try:
            doc = Document(file_stream)
            return "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            logger.error(f"Error al extraer texto de DOCX: {str(e)}")
            raise
            
    def _extract_text_with_fallback(self, content: bytes, file_type: str) -> str:
        """
        Intenta extraer texto usando métodos alternativos cuando el método principal falla.
        
        Args:
            content: Contenido binario del archivo
            file_type: Tipo de archivo ('pdf' o 'docx')
            
        Returns:
            str: Texto extraído
            
        Raises:
            ValueError: Si no se puede extraer el texto con ningún método
        """
        try:
            logger.info(f"Intentando método alternativo para extraer texto de {file_type}")
            
            if file_type.lower() == 'pdf':
                # Método alternativo para PDF usando PyPDF2
                try:
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                    text_parts = []
                    for page in pdf_reader.pages:
                        try:
                            text_parts.append(page.extract_text() or "")
                        except Exception as page_error:
                            logger.warning(f"Error al extraer texto de página PDF: {str(page_error)}")
                            continue
                    return "\n\n".join(text_parts)
                except Exception as pdf_error:
                    logger.error(f"Error en método alternativo PDF: {str(pdf_error)}")
                    
            elif file_type.lower() == 'docx':
                # Método alternativo para DOCX usando solo la biblioteca estándar
                try:
                    # Intentar con python-docx primero
                    doc = Document(io.BytesIO(content))
                    return "\n\n".join([p.text for p in doc.paragraphs])
                except Exception as docx_error:
                    logger.error(f"Error en método alternativo DOCX: {str(docx_error)}")
                    
            # Si llegamos aquí, todos los métodos fallaron
            raise ValueError(f"No se pudo extraer texto del archivo {file_type} con ningún método")
            
        except Exception as e:
            logger.error(f"Error en método alternativo de extracción: {str(e)}")
            raise ValueError(f"No se pudo extraer texto: {str(e)}")
    
    async def create(
        self, 
        obj_in: DocumentCreate, 
        user_id: int,
        **kwargs
    ) -> DocumentModel:
        """Crea un nuevo documento en la base de datos."""
        db_obj = DocumentModel(
            **obj_in.dict(exclude_unset=True),
            user_id=user_id,
            status=DocumentStatus.PENDING,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )
        
        self.db.add(db_obj)
        await self.db.commit()
        await self.db.refresh(db_obj)
        
        return db_obj
    
    async def update(
        self, 
        id: int, 
        obj_in: Union[DocumentUpdate, Dict[str, Any]],
        user_id: int,
        **kwargs
    ) -> Optional[DocumentModel]:
        """
        Actualiza un documento existente.
        
        Este método maneja la actualización de metadatos del documento y sus relaciones con bases de conocimiento.
        No se permite actualizar el archivo del documento directamente - para eso se debe usar upload_document.
        
        Args:
            id: ID del documento a actualizar
            obj_in: Datos de actualización (DocumentUpdate o diccionario)
            user_id: ID del usuario que realiza la actualización (debe ser el propietario)
            
        Returns:
            DocumentModel: El documento actualizado o None si no se encontró
            
        Raises:
            HTTPException: Si hay un error de validación o de base de datos
        """
        try:
            # Obtener el documento con bloqueo para actualización
            db_obj = await self.get(id, user_id=user_id, for_update=True)
            if not db_obj:
                return None
                
            # Convertir a diccionario si es un modelo Pydantic
            update_data = obj_in.dict(exclude_unset=True) if not isinstance(obj_in, dict) else obj_in.copy()
            
            # Validar que no se intente actualizar campos inmutables directamente
            immutable_fields = ['id', 'user_id', 'file_path', 'file_size', 'file_type', 'created_at']
            for field in immutable_fields:
                update_data.pop(field, None)
                
            # Actualizar metadatos
            update_data["updated_at"] = datetime.utcnow()
            
            # Manejar actualización de relaciones con knowledge bases si se proporciona
            knowledge_bases = update_data.pop('knowledge_bases', None)
            
            # Actualizar campos del documento
            for field, value in update_data.items():
                if hasattr(db_obj, field):
                    setattr(db_obj, field, value)
            
            # Actualizar relaciones con knowledge bases si se proporcionan
            if knowledge_bases is not None:
                # Eliminar relaciones existentes
                await self.db.execute(
                    delete(KnowledgeDocument).where(KnowledgeDocument.document_id == id)
                )
                
                # Crear nuevas relaciones
                for kb_data in knowledge_bases:
                    if isinstance(kb_data, dict):
                        kb_id = kb_data.get('knowledge_base_id')
                        kb_metadata = kb_data.get('metadata', {})
                    else:
                        kb_id = kb_data
                        kb_metadata = {}
                        
                    if kb_id:
                        # Verificar que la base de conocimiento exista y pertenezca al usuario
                        kb_exists = await self.db.execute(
                            exists(KnowledgeBase)
                            .where(
                                KnowledgeBase.id == kb_id,
                                KnowledgeBase.user_id == user_id
                            )
                            .select()
                        )
                        
                        if kb_exists.scalar():
                            kb_doc = KnowledgeDocument(
                                document_id=id,
                                knowledge_base_id=kb_id,
                                metadata=kb_metadata or {}
                            )
                            self.db.add(kb_doc)
            
            # Guardar cambios
            self.db.add(db_obj)
            await self.db.commit()
            await self.db.refresh(db_obj)
            
            # Cargar relaciones si se solicitan
            if kwargs.get('include_knowledge_bases', False):
                await self.db.refresh(db_obj, ['knowledge_bases'])
                
            return db_obj
            
        except HTTPException:
            await self.db.rollback()
            raise
            
        except Exception as e:
            await self.db.rollback()
            logger.error(f"Error al actualizar documento {id}: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error al actualizar el documento: {str(e)}"
            )
        
        return db_obj
    
    async def get(
        self, 
        document_id: int, 
        user_id: int, 
        include_knowledge_bases: bool = False,
        **kwargs
    ) -> Optional[DocumentModel]:
        """
        Obtiene un documento por su ID y el ID del usuario propietario.
        
        Args:
            document_id: ID del documento a recuperar
            user_id: ID del usuario propietario
            include_knowledge_bases: Si es True, incluye las relaciones con KnowledgeBases
            **kwargs: Argumentos adicionales:
                - for_update: Si es True, aplica un bloqueo de fila para actualización
                
        Returns:
            DocumentModel: El documento encontrado o None si no existe o no tiene permisos
                
        Raises:
            HTTPException: Si hay un error al acceder a la base de datos
        """
        try:
            # Construir la consulta base
            stmt = select(DocumentModel).where(
                DocumentModel.id == document_id,
                DocumentModel.user_id == user_id
            )
            
            # Aplicar bloqueo de fila si se solicita
            if kwargs.get('for_update', False):
                stmt = stmt.with_for_update()
                
            # Incluir relaciones con KnowledgeBases si se solicita
            if include_knowledge_bases:
                stmt = stmt.options(
                    selectinload(DocumentModel.knowledge_bases)
                    .selectinload(KnowledgeDocument.knowledge_base)
                )
                
            # Ejecutar la consulta
            result = await self.db.execute(stmt)
            document = result.scalar_one_or_none()
                
            return document
                
        except Exception as e:
            logger.error(f"Error al obtener documento {document_id}: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error al recuperar el documento: {str(e)}"
            )
    
    async def get_multi(
        self, 
        user_id: int, 
        skip: int = 0, 
        limit: int = 100,
        status: Optional[DocumentStatus] = None,
        doc_type: Optional[DocumentType] = None,
        knowledge_base_id: Optional[int] = None,
        include_knowledge_bases: bool = False,
        **kwargs
    ) -> List[DocumentModel]:
        """
        Obtiene múltiples documentos con opciones de filtrado avanzado.
        
        Permite filtrar documentos por estado, tipo y pertenencia a una base de conocimiento,
        así como incluir información detallada de las bases de conocimiento asociadas.
        
        Args:
            user_id: ID del usuario propietario
            skip: Número de documentos a omitir (paginación)
            limit: Número máximo de documentos a devolver (máx. 1000)
            status: Filtrar por estado del documento (opcional)
            doc_type: Filtrar por tipo de documento (opcional)
            knowledge_base_id: Filtrar por ID de base de conocimiento (opcional)
            include_knowledge_bases: Incluir información detallada de las bases de conocimiento
                
        Returns:
            List[DocumentModel]: Lista de documentos que coinciden con los criterios
                
        Raises:
            HTTPException: Si hay un error al acceder a la base de datos
        """
        try:
            # Validar parámetros
            if limit > 1000:
                limit = 1000  # Limitar el número máximo de resultados
                    
            # Construir la consulta base
            query = select(DocumentModel).where(DocumentModel.user_id == user_id)
                
            # Aplicar filtros opcionales
            if status is not None:
                query = query.where(DocumentModel.status == status)
                    
            if doc_type is not None:
                query = query.where(DocumentModel.doc_type == doc_type)
                
            # Si se especifica un knowledge_base_id, filtrar por él
            if knowledge_base_id is not None:
                # Subconsulta para obtener los IDs de documentos en la base de conocimiento
                subquery = select(KnowledgeDocument.document_id).where(
                    KnowledgeDocument.knowledge_base_id == knowledge_base_id
                ).scalar_subquery()
                    
                # Filtrar documentos que están en la base de conocimiento
                query = query.where(DocumentModel.id.in_(subquery))
                
            # Ordenar por fecha de creación descendente por defecto
            query = query.order_by(DocumentModel.created_at.desc())
                
            # Aplicar paginación
            query = query.offset(skip).limit(limit)
                
            # Cargar relaciones si es necesario
            if include_knowledge_bases:
                query = query.options(
                    selectinload(DocumentModel.knowledge_bases)
                    .selectinload(KnowledgeDocument.knowledge_base)
                )
                
            # Ejecutar la consulta
            result = await self.db.execute(query)
            documents = result.scalars().all()
            
            return documents
            
        except Exception as e:
            logger.error(f"Error al listar documentos: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error al recuperar los documentos: {str(e)}"
            )
            
    async def delete(self, id: int, user_id: int, **kwargs) -> bool:
        """
        Elimina un documento y sus relaciones asociadas.
        
        Este método elimina el documento de la base de datos, así como cualquier archivo físico
        asociado y sus relaciones con bases de conocimiento.
        
        Args:
            id: ID del documento a eliminar
            user_id: ID del usuario que realiza la eliminación (debe ser el propietario)
            
        Returns:
            bool: True si se eliminó correctamente, False si el documento no existe
            
        Raises:
            HTTPException: Si hay un error al eliminar el documento
        """
        try:
            # Obtener el documento con sus relaciones
            db_obj = await self.get(id, user_id=user_id, include_knowledge_bases=False)
            if not db_obj:
                return False
                
            # Eliminar el archivo físico si existe
            if db_obj.file_url and os.path.exists(db_obj.file_url):
                try:
                    os.remove(db_obj.file_url)
                    logger.info(f"Archivo físico eliminado: {db_obj.file_url}")
                except Exception as e:
                    logger.error(f"Error al eliminar el archivo {db_obj.file_url}: {str(e)}")
                    # No fallar si no se puede eliminar el archivo, solo registrar el error
            
            # Eliminar relaciones con bases de conocimiento primero
            stmt = delete(KnowledgeDocument).where(KnowledgeDocument.document_id == id)
            await self.db.execute(stmt)
            
            # Eliminar el documento de Weaviate si existe
            if db_obj.weaviate_id:
                try:
                    await weaviate_service.delete_document(db_obj.weaviate_id)
                    logger.info(f"Documento {id} eliminado de Weaviate")
                except Exception as e:
                    logger.error(f"Error al eliminar documento de Weaviate: {str(e)}")
                    # No fallar si hay un error con Weaviate
            
            # Eliminar el documento de la base de datos
            await self.db.delete(db_obj)
            await self.db.commit()
            
            logger.info(f"Documento {id} eliminado correctamente por el usuario {user_id}")
            return True
            
        except HTTPException:
            raise
            
        except Exception as e:
            await self.db.rollback()
            logger.error(f"Error al eliminar documento {id}: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error al eliminar el documento: {str(e)}"
            )
